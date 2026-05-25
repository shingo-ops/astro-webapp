"""
しんごさん向け Discord Bot セットアップ手順書（docx 生成スクリプト）

用途: docs/B-2_discord_setup_guide_for_shingo.docx を再生成する
実行場所: Mac 側（または python-docx が入った任意の環境）
使い方: PYENV_VERSION=3.11.1 python scripts/build_shingo_guide.py
        （または python -m pip install python-docx 済の任意の Python で）
前提: python-docx>=1.0 が pip install 済であること
出力: docs/B-2_discord_setup_guide_for_shingo.docx を上書き

2026-04-28 セッション成果として作成。
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# scripts/ の親（リポジトリルート）から docs/<出力> を組み立てる
_REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = str(_REPO_ROOT / "docs" / "B-2_discord_setup_guide_for_shingo.docx")


def add_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def style_run(run, *, size: int = 11, bold: bool = False, color: str | None = None, italic: bool = False) -> None:
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rFonts.set(qn(attr), "Yu Gothic")


def add_para(doc, text: str, *, size: int = 11, bold: bool = False, color: str | None = None,
             align=None, space_after: int = 6, indent: float = 0.0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text: str, *, level: int):
    sizes = {0: 22, 1: 18, 2: 14, 3: 12}
    colors = {0: "1F2937", 1: "1E40AF", 2: "047857", 3: "374151"}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level >= 1 else 0)
    p.paragraph_format.space_after = Pt(6)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    style_run(run, size=sizes.get(level, 11), bold=True, color=colors.get(level, "000000"))
    return p


def add_bullet(doc, text: str, *, level: int = 0, size: int = 11, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    if bold_prefix:
        run_prefix = p.add_run(bold_prefix)
        style_run(run_prefix, size=size, bold=True)
        run = p.add_run(text)
    else:
        run = p.add_run(text)
    style_run(run, size=size, bold=False)
    return p


def add_callout(doc, label: str, body: str, *, color: str = "FEF3C7", text_color: str = "92400E"):
    """色付きの「ひとことメモ」枠を 1 行のテーブルで作る。"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    add_shading(cell, color)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)
    run_label = para.add_run(f"{label}  ")
    style_run(run_label, size=10, bold=True, color=text_color)
    run_body = para.add_run(body)
    style_run(run_body, size=10, color="111827")
    # 横幅を本文と揃える
    table.autofit = False
    table.columns[0].width = Cm(16)
    cell.width = Cm(16)
    return table


def add_table_with_header(doc, headers: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            table.columns[i].width = Cm(w)
    # ヘッダー
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        add_shading(cell, "1E40AF")
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        style_run(run, size=10, bold=True, color="FFFFFF")
        if col_widths_cm:
            cell.width = Cm(col_widths_cm[i])
    # データ
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            style_run(run, size=10)
            if col_widths_cm:
                cell.width = Cm(col_widths_cm[c_idx])
    # 余白
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return table


def add_step(doc, n: int, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.6)
    run_n = p.add_run(f"Step {n}.  ")
    style_run(run_n, size=11, bold=True, color="1E40AF")
    run = p.add_run(text)
    style_run(run, size=11)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "9CA3AF")
    bdr.append(bottom)
    pPr.append(bdr)


# ===== Build document =====
doc = Document()

# 余白を狭めに
for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# タイトルページ
add_heading(doc, "Discord Bot セットアップ手順書", level=0)
add_para(doc, "Sales Anchor B-2: Discord 連携 M1 タスク（しんごさん作業分）",
         size=12, color="6B7280", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, "2026-04-28 作成 / 開発パートナー: hitoshi（Claude Code）",
         size=10, color="9CA3AF", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

# はじめに
add_heading(doc, "はじめに", level=1)
add_para(doc,
         "Sales Anchor の Discord 連携機能（リードからの DM 受信 + 自動返信）を本番稼働させる"
         "ため、しんごさんに 4 つのセットアップ作業をお願いします。所要時間は合計 30〜60 分程度です。",
         size=11)
add_callout(doc, "💡 たとえ話",
            "これは「お店に新人スタッフ（配達員）を採用する」作業によく似ています。Bot は Sales Anchor 専用の "
            "配達員、Discord 開発者ポータルは採用面接の窓口、Bot Token は社員証 + 倉庫の鍵、"
            "パスワードマネージャー（Bitwarden / 1Password など）はその鍵を保管する金庫、"
            "テスト用ギルドは新人研修用の練習スペース、というイメージです。",
            color="EFF6FF", text_color="1E3A8A")

# 用語マッピング表
add_heading(doc, "用語マッピング（たとえ話 ↔ 実物）", level=2)
add_table_with_header(
    doc,
    ["お店のたとえ", "Discord での実物", "本書の登場場所"],
    [
        ["新人配達員", "Discord Bot（アプリ + Bot 機能）", "タスク 1"],
        ["配達員に「家の中まで入っていい権限」を付与", "Privileged Intents（MESSAGE CONTENT / SERVER MEMBERS）", "タスク 2"],
        ["社員証 + 倉庫の鍵（盗まれたら危ない）", "Bot Token（再発行はできるが流出は事故）", "タスク 3"],
        ["金庫", "パスワードマネージャー（Bitwarden / 1Password 等）", "タスク 3"],
        ["新人研修用の練習スペース", "テスト用 Discord ギルド（本番と別）", "タスク 4"],
        ["お店の本店", "HIGH LIFE JPN の本番 Discord ギルド", "後日（Bot 動作確認後）"],
    ],
    col_widths_cm=[5.5, 6.0, 4.5],
)

# タスク一覧
add_heading(doc, "タスク一覧", level=1)
add_table_with_header(
    doc,
    ["#", "タスク", "所要時間", "難易度"],
    [
        ["1", "Discord 開発者ポータルで Bot を作成", "10 分", "★☆☆"],
        ["2", "Privileged Intents を ON にする", "3 分", "★☆☆"],
        ["3", "Bot Token をパスワードマネージャーに保管", "5 分", "★★☆"],
        ["4", "テスト用 Discord ギルドを準備", "10 分", "★☆☆"],
    ],
    col_widths_cm=[1.0, 9.0, 3.0, 3.0],
)

add_horizontal_rule(doc)

# ============================================================
# Task 1
# ============================================================
add_heading(doc, "タスク 1: Discord 開発者ポータルで Bot を作成", level=1)

add_callout(doc, "🍱 たとえ話",
            "新人配達員を雇うため「採用面接」を Discord に申請する作業。"
            "ここで作るのは「Bot 本体」ではなく「アプリ枠」で、"
            "面接が受かったらアプリ枠の中で Bot 機能を有効化する流れになります。")

add_heading(doc, "なぜ必要か", level=3)
add_para(doc,
         "Discord は「誰が運用している Bot か」を必ず特定するため、最初に開発者として "
         "アプリを登録する必要があります。Sales Anchor はテナント（HIGH LIFE JPN）ごとに "
         "別 Bot を運用する方針（per-tenant Bot）なので、HIGH LIFE JPN 専用の枠を 1 つ作ります。")

add_heading(doc, "やり方", level=3)
add_step(doc, 1, "Discord 開発者ポータルにアクセス: https://discord.com/developers/applications")
add_step(doc, 2, "しんごさんの個人 Discord アカウント（Treasure Island JP のメールに紐づくもの）でログイン。"
                  "未作成なら新規登録してください。")
add_step(doc, 3, "右上の「New Application」をクリック")
add_step(doc, 4, "アプリ名を入力。例: Sales Anchor (HIGH LIFE JPN)")
add_step(doc, 5, "規約に同意 → Create で作成")
add_step(doc, 6, "左メニューの「Bot」をクリック → 「Add Bot」（または「Reset Token」が出ていれば既に Bot 化済み）")

add_callout(doc, "✅ 完了の目安",
            "アプリ画面の上部に Application ID（数字 18 桁）が表示され、左メニューに「Bot」項目が "
            "あればここまで OK です。",
            color="ECFDF5", text_color="065F46")

add_horizontal_rule(doc)

# ============================================================
# Task 2
# ============================================================
add_heading(doc, "タスク 2: Privileged Intents を ON にする", level=1)

add_callout(doc, "🔑 たとえ話",
            "新人配達員に「お客様の家の中まで荷物を運んでいいよ」と特別な権限を与える作業。"
            "通常の Bot は『メッセージが届いた事実』しか分からず、中身（本文）を読むには Discord に "
            "別途許可を申請する必要があります。これは Discord 側のプライバシー保護のための仕組みです。")

add_heading(doc, "やり方", level=3)
add_step(doc, 1, "タスク 1 で開いた Bot 設定画面（左メニュー「Bot」）に戻る")
add_step(doc, 2, "下にスクロールして「Privileged Gateway Intents」のセクションを見つける")
add_step(doc, 3, "次の 2 つのトグルを ON にする:")
add_bullet(doc, "MESSAGE CONTENT INTENT  ←メッセージ本文を読む権限",
            level=1, bold_prefix="① ")
add_bullet(doc, "SERVER MEMBERS INTENT  ←サーバー参加者の一覧を取得する権限",
            level=1, bold_prefix="② ")
add_step(doc, 4, "画面下の「Save Changes」をクリックして保存")

add_callout(doc, "⚠️ 重要",
            "ここを ON にし忘れると、Bot を起動しても READY イベントが出ず接続できません。"
            "Discord ガイドラインで 100 ギルド以下なら申請不要ですが、トグル自体は guild 数に "
            "関係なく必須なのでお忘れなく。",
            color="FEF2F2", text_color="991B1B")

add_horizontal_rule(doc)

# ============================================================
# Task 3
# ============================================================
add_heading(doc, "タスク 3: Bot Token をパスワードマネージャーに保管", level=1)

add_callout(doc, "📌 前提",
            "本書では「パスワードマネージャー」と一般化して書きます。具体的には Bitwarden / 1Password / "
            "他社製品（KeePass, Dashlane 等）のいずれでも構いません。社内で利用予定のツールに読み替えて"
            "進めてください。",
            color="F3F4F6", text_color="374151")

add_callout(doc, "🔐 たとえ話",
            "新人配達員の「社員証 + 倉庫の鍵」を金庫にしまう作業。"
            "鍵を盗まれるとお店全体の在庫が危険にさらされるので、画面に表示された瞬間に "
            "金庫に運ぶ感覚で手早く保管します。")

add_heading(doc, "なぜ最重要か", level=3)
add_para(doc,
         "Bot Token は Bot の「身代わり」で、これさえあれば誰でも HIGH LIFE JPN の Discord "
         "メッセージを送受信できてしまいます。流出すると Bot を「乗っ取られた」状態となり、"
         "再発行（Token Reset）まで第三者から悪用されます。")
add_para(doc,
         "Discord は Token を一度しか画面に表示しません（タブを閉じると二度と取れない）。"
         "再表示はできず、「Reset Token」で新しい Token を発行 → 古い Token は失効、という運用になります。",
         size=11)

add_heading(doc, "やり方", level=3)
add_step(doc, 1, "Bot 設定画面の上部にある「Reset Token」をクリック")
add_step(doc, 2, "「2 段階認証コード」が要求されたら、しんごさんの認証アプリで生成して入力")
add_step(doc, 3, "新しい Token が表示される（数十文字の英数字+ドット）")
add_step(doc, 4, "★その場でパスワードマネージャーを開いて新規エントリを作成★")
add_step(doc, 5, "エントリ名: 例: Discord Bot Token - Sales Anchor (HIGH LIFE JPN)")
add_step(doc, 6, "保存項目（推奨）:")
add_bullet(doc, "Token 本文 → Password 欄に貼る（長い文字列）", level=1)
add_bullet(doc, "Application ID → Username 欄 or Notes に貼る（タスク 1 で確認した 18 桁）", level=1)
add_bullet(doc, "発行日 → Notes / Custom Field（例: 2026-04-28）", level=1)
add_bullet(doc, "Reset 履歴 → Notes（例: 2026-04-28 初回発行）", level=1)
add_step(doc, 7, "パスワードマネージャー側で保存完了を確認したらブラウザのタブを閉じて OK")

add_callout(doc, "💡 ツール別の補足",
            "Bitwarden: 「Login 項目を追加」→ Name / Username / Password / Notes に上記を入れる。\n"
            "1Password: 「Password」テンプレで作成、または「API Credential」テンプレが Token 用に最適。\n"
            "両方とも Custom Field（カスタムフィールド）が使えるので Application ID / 発行日はそこに入れると見やすい。",
            color="F0FDF4", text_color="166534")

add_heading(doc, "やってはいけないこと（NG 行動）", level=3)
add_table_with_header(
    doc,
    ["NG 行動", "なぜ NG か"],
    [
        ["メールやチャット（Slack 等）に Token を貼り付ける", "送信履歴・受信箱に残り、アカウント侵害時に芋づる式に流出"],
        ["スクリーンショットを撮ってクラウドに保存", "Google Photos / iCloud に同期され検索可能に"],
        [".env / コードリポジトリにコミット", "GitHub 公開した瞬間に Bot Token がスキャンされ Bot 乗っ取り"],
        ["「これあとで貼ろう」とメモに残す", "貼り忘れて画面を閉じ、二度と取れず Reset → 再発行になる"],
    ],
    col_widths_cm=[6.5, 9.5],
)

add_callout(doc, "🛟 もし画面を閉じてしまったら",
            "焦らず大丈夫です。「Reset Token」をもう一度クリックすれば新しい Token が発行されます。"
            "古い Token は失効するので「漏れた可能性」も自動的に消えます（運用的には安全側）。",
            color="ECFDF5", text_color="065F46")

add_horizontal_rule(doc)

# ============================================================
# Task 4
# ============================================================
add_heading(doc, "タスク 4: テスト用 Discord ギルドを準備", level=1)

add_callout(doc, "🏠 たとえ話",
            "新人配達員を本番（HIGH LIFE JPN サーバー）に投入する前の「研修用の家」を用意する作業。"
            "本番ギルドにはお客様や既存 Bot（Tickets Bot / John-Bot）がすでに住んでいるので、"
            "いきなり Jarvis Bot を入れると役割が衝突する可能性があります。"
            "まずは誰もいない研修ハウスで動作確認してから本番に移します。")

add_heading(doc, "やり方", level=3)
add_step(doc, 1, "しんごさん個人の Discord アプリで「サーバーを追加」（左メニュー一番下の「+」）")
add_step(doc, 2, "「オリジナルの作成」→「自分と友達のため」を選択")
add_step(doc, 3, "サーバー名: Sales Anchor Test など分かりやすい名前で OK")
add_step(doc, 4, "Bot を招待するための招待リンクを作成:")
add_bullet(doc, "Discord 開発者ポータル → 該当アプリ → 左メニュー「OAuth2」→ 「URL Generator」", level=1)
add_bullet(doc, "Scopes で  bot  にチェック", level=1)
add_bullet(doc, "Bot Permissions で次にチェック: Send Messages / Read Messages / Read Message History", level=1)
add_bullet(doc, "下部に生成された URL をコピー", level=1)
add_step(doc, 5, "コピーした URL をブラウザで開く → 招待先に「Sales Anchor Test」サーバーを選択 → 認証")
add_step(doc, 6, "Discord アプリでテストサーバーを開き、メンバー一覧に Bot が表示されていれば成功")

add_callout(doc, "✅ 完了の目安",
            "テストサーバーのメンバー一覧に Bot が「オフライン」状態で表示されていれば OK。"
            "オンラインに切り替わるのは hitoshi 側で Token を VPS にセットして起動した後です。",
            color="ECFDF5", text_color="065F46")

add_horizontal_rule(doc)

# ============================================================
# After all tasks
# ============================================================
add_heading(doc, "全タスク完了後の流れ", level=1)
add_para(doc,
         "しんごさんから次の 2 つの完了連絡をいただいたら、私（hitoshi）の方で残作業を進めます:",
         size=11)
add_bullet(doc, "パスワードマネージャーに Token を保管完了", bold_prefix="① ")
add_bullet(doc, "テスト用ギルドに Bot を招待完了", bold_prefix="② ")

add_para(doc, "私の方で実施する内容:", size=11, bold=True)
add_bullet(doc, "VPS の環境変数 DISCORD_BOT_TOKEN_4 に Token を設定")
add_bullet(doc, "discord-gateway コンテナを再起動")
add_bullet(doc, "READY イベントログを確認（Bot が Discord に接続できたか）")
add_bullet(doc, "コンテナ再起動 → session resume が動くか検証")
add_bullet(doc, "5 分以上の安定接続を確認")
add_para(doc,
         "ここまで完了で B-2「Discord Gateway 常駐プロセス」の M2（Skeleton 段階）動作検証完了となります。"
         "M3（受信メッセージを Sales Anchor に取り込む処理）は別セッションで実装します。",
         size=11)

add_horizontal_rule(doc)

# トラブルシューティング
add_heading(doc, "困ったとき（トラブルシューティング）", level=1)
add_table_with_header(
    doc,
    ["症状", "考えられる原因", "対処"],
    [
        ["Discord にログインできない",
         "Treasure Island JP のメールに紐づく Discord アカウントが未作成",
         "新規登録（無料、メール認証のみ）"],
        ["「New Application」が押せない",
         "Discord 利用規約 / 開発者規約への同意が未完了",
         "ポップアップに同意してリトライ"],
        ["Token をパスワードマネージャーに保存し忘れて画面を閉じた",
         "Discord は Token を 1 度しか表示しない",
         "「Reset Token」で再発行（古い Token は失効するので安全）"],
        ["Privileged Intents の「Save」ボタンが押せない",
         "未保存の他項目があるか、規約画面が出ている",
         "ページをスクロールして警告メッセージを確認、規約があれば同意"],
        ["Bot 招待 URL を開いてもサーバーが選べない",
         "URL の Scope に bot が含まれていない",
         "URL Generator で再生成（必須: bot, applications.commands は任意）"],
        ["招待後もテストサーバーに Bot が出てこない",
         "招待時のサーバー選択を間違えた / 権限不足",
         "テストサーバーから一度キックして招待 URL を再度開く"],
    ],
    col_widths_cm=[4.5, 5.5, 6.0],
)

add_horizontal_rule(doc)

# 付録
add_heading(doc, "付録: 用語の補足説明", level=1)
add_table_with_header(
    doc,
    ["用語", "ざっくり意味"],
    [
        ["Discord Bot",
         "Discord 上で人間の代わりに動く自動アカウント。Jarvis Bot は「メッセージの送受信」のみ担当します"],
        ["Application / Bot",
         "Discord ではまず Application（アプリ枠）を作り、その中で Bot 機能を有効化する 2 段階の構造"],
        ["Token",
         "Bot を「これは本物です」と Discord に証明するための長い文字列。流出 = Bot 乗っ取り"],
        ["Intent",
         "Bot が Discord から受け取るイベントの種類を絞り込む仕組み（最小権限の原則）"],
        ["Privileged Intent",
         "プライバシー上重要な情報（メッセージ本文、メンバー一覧など）を取得するための特別な権限"],
        ["Gateway",
         "Discord と Bot をつなぐ常時接続の通信路（WebSocket）。今回追加したのはこのプロセス"],
        ["Guild",
         "Discord 内部用語で「サーバー」のこと。HIGH LIFE JPN のサーバー = HIGH LIFE JPN ギルド"],
        ["session resume",
         "ネットワーク切断などで Bot が落ちても、再接続時に過去の状態を引き継ぐ仕組み"],
        ["per-tenant Bot",
         "テナント（HIGH LIFE JPN）ごとに別の Bot を運用する方式。漏洩時の影響範囲を限定するため"],
    ],
    col_widths_cm=[4.0, 12.0],
)

add_horizontal_rule(doc)

# 連絡先
add_heading(doc, "ご質問はお気軽に", level=1)
add_para(doc,
         "本書の手順で詰まった箇所、用語の意味が不明な箇所、UI が変わっていて見つからない箇所など、"
         "どんなことでも遠慮なく hitoshi（Claude Code）にお声がけください。",
         size=11)
add_para(doc,
         "Discord は UI 変更が頻繁なので、最新のスクリーンショットを送っていただけると素早く対応できます。",
         size=11)

# 保存
import os
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f"OK: saved to {OUT_PATH}")
print(f"   sections: {len(doc.sections)}, paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}")
